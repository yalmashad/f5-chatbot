import os
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

import requests
import streamlit as st
from docx import Document
from dotenv import dotenv_values, load_dotenv
from openai import OpenAI
from pypdf import PdfReader

from app_settings import (
    DEFAULT_COMPATIBLE_BASE_URL,
    DEFAULT_COMPATIBLE_MODEL,
    DEFAULT_GUARDRAIL_HOSTNAME,
    DEFAULT_OLLAMA_BASE_URL,
    DEFAULT_OLLAMA_MODEL,
    DEFAULT_OPENAI_MODEL,
    get_available_model_providers,
    get_selected_model,
    reset_session_settings,
    seed_session_settings,
    update_session_settings,
)


BASE_DIR = Path(__file__).resolve().parent
ENV_PATH = BASE_DIR / ".env"
load_dotenv(dotenv_path=ENV_PATH, override=True)

COMMON_OPENAI_MODELS = (
    "gpt-5.5",
    "gpt-5.4",
    "gpt-5.4-mini",
    "gpt-4.1",
    "gpt-4o",
)
OTHER_MODEL_OPTION = "Other"
MAX_DOCUMENT_FILE_BYTES = 10 * 1024 * 1024
MAX_EXTRACTED_TEXT_CHARS = 100_000
SUPPORTED_DOCUMENT_EXTENSIONS = (".pdf", ".docx")


class DocumentInspectionError(Exception):
    """Raised when a document cannot be safely inspected."""


@dataclass(frozen=True)
class ExtractedDocument:
    filename: str
    extension: str
    size_bytes: int
    text: str

    @property
    def char_count(self) -> int:
        return len(self.text)


def validate_document_upload(uploaded_file) -> str:
    filename = uploaded_file.name or ""
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_DOCUMENT_EXTENSIONS:
        raise DocumentInspectionError("Only PDF and DOCX documents are supported.")

    size_bytes = getattr(uploaded_file, "size", 0) or 0
    if size_bytes > MAX_DOCUMENT_FILE_BYTES:
        raise DocumentInspectionError(
            "Document is too large to inspect safely. Maximum size is 10 MB."
        )

    return extension


def normalize_extracted_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def escape_markdown_text(text: str) -> str:
    return text.translate(str.maketrans({char: f"\\{char}" for char in r"\\`*_{}[]()#+-.!|>"}))


def format_attachment_caption(filename: str) -> str:
    return f"Attached: {escape_markdown_text(filename)}"


def extract_pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    page_text = []
    for page in reader.pages:
        page_text.append(page.extract_text() or "")
        normalized_text = normalize_extracted_text("\n".join(page_text))
        if len(normalized_text) > MAX_EXTRACTED_TEXT_CHARS:
            raise DocumentInspectionError("Document text is too large to inspect safely.")
    return normalize_extracted_text("\n".join(page_text))


def extract_docx_text(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))
    paragraphs = []
    for paragraph in document.paragraphs:
        paragraphs.append(paragraph.text)
        normalized_text = normalize_extracted_text("\n".join(paragraphs))
        if len(normalized_text) > MAX_EXTRACTED_TEXT_CHARS:
            raise DocumentInspectionError("Document text is too large to inspect safely.")
    return normalize_extracted_text("\n".join(paragraphs))


def extract_uploaded_document(uploaded_file) -> ExtractedDocument:
    extension = validate_document_upload(uploaded_file)
    file_bytes = uploaded_file.getvalue()
    if len(file_bytes) > MAX_DOCUMENT_FILE_BYTES:
        raise DocumentInspectionError(
            "Document is too large to inspect safely. Maximum size is 10 MB."
        )

    if extension == ".pdf":
        text = extract_pdf_text(file_bytes)
    elif extension == ".docx":
        text = extract_docx_text(file_bytes)
    else:
        raise DocumentInspectionError("Only PDF and DOCX documents are supported.")

    build_document_inspection_payload("", uploaded_file.name, text)
    return ExtractedDocument(
        filename=uploaded_file.name,
        extension=extension,
        size_bytes=len(file_bytes),
        text=text,
    )


def build_document_inspection_payload(prompt: str, filename: str, document_text: str) -> str:
    normalized_text = normalize_extracted_text(document_text)
    if not normalized_text:
        raise DocumentInspectionError("No inspectable text was found in the document.")
    if len(normalized_text) > MAX_EXTRACTED_TEXT_CHARS:
        raise DocumentInspectionError("Document text is too large to inspect safely.")

    return (
        f"User prompt:\n{prompt}\n\n"
        f"Document name:\n{filename}\n\n"
        "Extracted document text:\n"
        f"{normalized_text}"
    )


def build_document_model_messages(prompt: str, filename: str, document_text: str) -> list[dict[str, str]]:
    normalized_text = normalize_extracted_text(document_text)
    return [
        {"role": "system", "content": "You are a helpful assistant."},
        {
            "role": "user",
            "content": (
                f"User prompt:\n{prompt}\n\n"
                f"Document name:\n{filename}\n\n"
                "Extracted document text:\n"
                f"{normalized_text}"
            ),
        },
    ]


def build_model_response_scan_payload(response_text: str) -> str:
    return f"Model response:\n{response_text}"


def redact_sensitive_debug_data(value, sensitive_values: list[str] | None):
    redaction = "[redacted document content]"
    normalized_sensitive_values = [item for item in (sensitive_values or []) if item]

    if isinstance(value, dict):
        return {
            key: redact_sensitive_debug_data(item, normalized_sensitive_values)
            for key, item in value.items()
        }
    if isinstance(value, list):
        return [
            redact_sensitive_debug_data(item, normalized_sensitive_values)
            for item in value
        ]
    if isinstance(value, tuple):
        return tuple(
            redact_sensitive_debug_data(item, normalized_sensitive_values)
            for item in value
        )
    if isinstance(value, str):
        redacted_value = value
        for sensitive_value in normalized_sensitive_values:
            redacted_value = redacted_value.replace(sensitive_value, redaction)
        return redacted_value
    return value


def document_metadata(extracted_document: ExtractedDocument | None) -> dict | None:
    if extracted_document is None:
        return None
    return {
        "filename": extracted_document.filename,
        "extension": extracted_document.extension,
        "size_bytes": extracted_document.size_bytes,
        "char_count": extracted_document.char_count,
    }


def render_debug_details(message: dict) -> None:
    if "cai_json" in message and message["cai_json"] is not None:
        with st.expander("F5 Guardrail JSON"):
            st.json(message["cai_json"])
    if "scan_in" in message and message["scan_in"] is not None:
        with st.expander("Guardrail scan - input JSON"):
            st.json(message["scan_in"])
    if "scan_out" in message and message["scan_out"] is not None:
        with st.expander("Guardrail scan - output JSON"):
            st.json(message["scan_out"])
    if "document" in message and message["document"] is not None:
        with st.expander("Document extraction metadata"):
            st.json(message["document"])


def require_env(var_name: str, var_value: str | None) -> None:
    if not var_value:
        raise RuntimeError(f"Missing {var_name}. Enter it in Settings or provide it in .env.")


def get_openai_client(api_key: str) -> OpenAI:
    require_env("OPENAI_API_KEY", api_key)
    return OpenAI(api_key=api_key)


def get_openai_compatible_client(base_url: str, api_key: str) -> OpenAI:
    require_env("OPENAI_COMPATIBLE_BASE_URL", base_url)
    require_env("OPENAI_COMPATIBLE_API_KEY", api_key)
    return OpenAI(base_url=base_url, api_key=api_key)


def get_ollama_client(base_url: str) -> OpenAI:
    require_env("OLLAMA_BASE_URL", base_url)
    return OpenAI(base_url=base_url, api_key="ollama")


def get_llm_client(settings: dict[str, str]) -> OpenAI:
    if settings["model_provider"] == "Ollama":
        return get_ollama_client(settings["ollama_base_url"])
    if settings["model_provider"] == "OpenAI compatible":
        return get_openai_compatible_client(
            settings["openai_compatible_base_url"],
            settings["openai_compatible_api_key"],
        )
    return get_openai_client(settings["openai_api_key"])


def get_ollama_tags_url(base_url: str) -> str:
    normalized = base_url.rstrip("/")
    if normalized.endswith("/v1"):
        normalized = normalized[:-3]
    return f"{normalized}/api/tags"


def fetch_ollama_models(base_url: str) -> tuple[list[str], str | None]:
    try:
        resp = requests.get(get_ollama_tags_url(base_url), timeout=5)
        resp.raise_for_status()
        data = resp.json()
        models = sorted(
            {model.get("name", "") for model in data.get("models", []) if model.get("name")}
        )
        return models, None
    except Exception as e:
        return [], str(e)


def get_model_error_hint(exc: Exception, settings: dict[str, str]) -> str | None:
    message = str(exc)
    if settings["model_provider"] == "OpenAI" and (
        "model_not_found" in message or "does not have access to model" in message
    ):
        return (
            f"Selected OpenAI model `{get_selected_model(settings)}` is not available for this "
            "API key or project. Choose a different OpenAI model in Settings or switch to Ollama."
        )
    if settings["model_provider"] == "OpenAI compatible" and (
        "Connection error" in message or "NotFoundError" in message or "404" in message
    ):
        return (
            "The OpenAI-compatible endpoint did not accept the request. Confirm the base URL "
            "includes the API version path if required, for example `https://host.example.com/v1`."
        )
    return None


def cai_scanapi(
    text: str,
    api_key: str,
    scan_url: str,
) -> tuple[bool, dict | None]:
    """
    F5 Guardrail Scan API: returns (cleared?, full_json_or_none).
    """
    require_env("GUARDRAIL_API_KEY", api_key)

    try:
        resp = requests.post(
            scan_url,
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
                "Accept": "application/json",
            },
            json={"input": text},
            timeout=60,
            allow_redirects=False,
        )

        if not resp.ok:
            return False, {"http_status": resp.status_code, "body": resp.text}

        data = resp.json()
        result = data.get("result")
        if not result:
            return False, {"error": "Missing 'result' in response", "data": data}

        outcome = result.get("outcome", "unknown")
        return outcome == "cleared", data

    except Exception as e:
        return False, {"error": str(e)}


def cai_promptapi(prompt: str, api_key: str, prompt_api_url: str) -> tuple[str | None, dict]:
    """
    F5 Guardrail Prompt API (Inline): returns (assistant_text_or_none, full_json).
    """
    require_env("GUARDRAIL_API_KEY", api_key)

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "Accept": "application/json",
    }

    body = {"input": prompt}

    resp = requests.post(
        prompt_api_url,
        headers=headers,
        json=body,
        timeout=60,
        allow_redirects=False,
    )

    if not resp.ok:
        return None, {"http_status": resp.status_code, "body": resp.text}

    data = resp.json()
    result = data.get("result")
    if not result:
        return None, {"error": "Missing 'result' in response", "data": data}

    outcome = result.get("outcome", "unknown")
    if outcome != "cleared":
        return None, data

    return result.get("response", ""), data


def llm_chat(
    prompt: str,
    settings: dict[str, str],
    messages: list[dict[str, str]] | None = None,
) -> str:
    """
    Chat completion using either OpenAI or a local Ollama endpoint.
    """
    client = get_llm_client(settings)
    model = get_selected_model(settings)
    chat_messages = messages or [
        {"role": "system", "content": "You are a helpful assistant."},
        {"role": "user", "content": prompt},
    ]

    completion = client.chat.completions.create(
        model=model,
        messages=chat_messages,
    )
    return completion.choices[0].message.content or ""


st.set_page_config(page_title="Secure Chatbot", layout="centered")
st.title("🛡️ F5 Secure Chatbot")

env_values = {**os.environ, **dotenv_values(ENV_PATH)}
settings = seed_session_settings(st.session_state, env_values)

with st.sidebar:
    st.header("Guardrail settings")

    guardrail_enabled = st.checkbox("Enable Guardrail", value=True)

    if guardrail_enabled:
        guardrail_mode = st.selectbox(
            "Mode",
            ["Inline", "Out-of-band"],
            index=0,
            help="Inline uses the F5 Guardrail Prompt API. Out-of-band scans before and after the model call.",
        )
    else:
        guardrail_mode = "Disabled"

    provider_settings_disabled = guardrail_enabled and guardrail_mode == "Inline"
    if provider_settings_disabled:
        st.info(
            "Inline mode uses the model configured in F5 Guardrail, so model settings here do not apply."
        )

    show_debug = st.checkbox("Show debug details", value=False)
    if st.button("Clear chat"):
        st.session_state.messages = []
        st.session_state.last_mode = (
            guardrail_enabled,
            guardrail_mode,
            settings["model_provider"],
            get_selected_model(settings),
        )
        st.rerun()

    st.divider()

    with st.expander("Settings", expanded=True):
        available_model_providers = get_available_model_providers(settings)
        provider_index = available_model_providers.index(settings["model_provider"])
        model_provider = st.selectbox(
            "Model provider",
            available_model_providers,
            index=provider_index,
            disabled=provider_settings_disabled,
            key="settings_model_provider",
        )

        with st.form("settings_form", enter_to_submit=False):
            openai_model = settings["openai_model"]
            openai_api_key = settings["openai_api_key"]
            openai_compatible_model = settings["openai_compatible_model"]
            openai_compatible_base_url = settings["openai_compatible_base_url"]
            openai_compatible_api_key = settings["openai_compatible_api_key"]
            ollama_model = settings["ollama_model"]
            ollama_base_url = settings["ollama_base_url"]
            guardrail_api_key = settings["guardrail_api_key"]
            guardrail_hostname = settings["guardrail_hostname"]

            with st.container(border=True):
                st.subheader("Model Provider")
                if model_provider == "OpenAI":
                    openai_model_options = list(COMMON_OPENAI_MODELS) + [OTHER_MODEL_OPTION]
                    use_custom_openai = settings["openai_model"] not in openai_model_options
                    if use_custom_openai:
                        openai_model_index = len(openai_model_options) - 1
                    else:
                        openai_model_index = openai_model_options.index(settings["openai_model"])

                    selected_openai_model = st.selectbox(
                        "OpenAI model",
                        openai_model_options,
                        index=openai_model_index,
                        help="Choose a common OpenAI model or enter one manually.",
                        disabled=provider_settings_disabled,
                    )

                    if selected_openai_model == OTHER_MODEL_OPTION:
                        openai_model = st.text_input(
                            "Other OpenAI model",
                            value=settings["openai_model"] if use_custom_openai else "",
                            help="Use this if the model you want is not listed.",
                            disabled=provider_settings_disabled,
                        )
                    else:
                        openai_model = selected_openai_model

                    openai_api_key = st.text_input(
                        "OpenAI API key",
                        value="",
                        placeholder=(
                            "Configured for this session"
                            if settings["openai_api_key"]
                            else ""
                        ),
                        type="password",
                        help="Leave blank to keep the current session value.",
                        disabled=provider_settings_disabled,
                    )
                elif model_provider == "OpenAI compatible":
                    openai_compatible_base_url = st.text_input(
                        "OpenAI-compatible base URL",
                        value=settings["openai_compatible_base_url"],
                        placeholder="https://api.example.com/v1",
                        help="Use the provider's OpenAI-compatible API base URL, usually ending in /v1.",
                        disabled=provider_settings_disabled,
                    )
                    openai_compatible_model = st.text_input(
                        "OpenAI-compatible model",
                        value=settings["openai_compatible_model"],
                        placeholder=DEFAULT_COMPATIBLE_MODEL,
                        disabled=provider_settings_disabled,
                    )
                    openai_compatible_api_key = st.text_input(
                        "OpenAI-compatible API key",
                        value="",
                        placeholder=(
                            "Configured for this session"
                            if settings["openai_compatible_api_key"]
                            else ""
                        ),
                        type="password",
                        help="Leave blank to keep the current session value.",
                        disabled=provider_settings_disabled,
                    )
                else:
                    ollama_base_url = st.text_input(
                        "Ollama base URL",
                        value=settings["ollama_base_url"],
                        help="Usually http://localhost:11434/v1",
                        disabled=provider_settings_disabled,
                    )

                    ollama_models, ollama_models_error = fetch_ollama_models(ollama_base_url)
                    if ollama_models:
                        ollama_options = list(ollama_models)
                        use_custom_ollama = settings["ollama_model"] not in ollama_options
                        if use_custom_ollama:
                            ollama_options.append("Custom...")
                            default_index = len(ollama_options) - 1
                        else:
                            default_index = ollama_options.index(settings["ollama_model"])

                        selected_ollama_model = st.selectbox(
                            "Ollama model",
                            ollama_options,
                            index=default_index,
                            disabled=provider_settings_disabled,
                        )

                        if selected_ollama_model == "Custom...":
                            ollama_model = st.text_input(
                                "Custom Ollama model",
                                value=settings["ollama_model"] if use_custom_ollama else "",
                                help="Use this if your local model is not listed yet.",
                                disabled=provider_settings_disabled,
                            )
                        else:
                            ollama_model = selected_ollama_model
                    else:
                        ollama_model = st.text_input(
                            "Ollama model",
                            value=settings["ollama_model"],
                            help="Example: llama3.2, mistral, qwen2.5",
                            disabled=provider_settings_disabled,
                        )
                        ollama_models_error = (
                            f"Could not load Ollama models from `{get_ollama_tags_url(ollama_base_url)}`. "
                            "You can still enter a model name manually."
                        )

                    if ollama_models_error:
                        st.caption(ollama_models_error)

            with st.container(border=True):
                st.subheader("F5 Guardrail")
                guardrail_hostname = st.text_input(
                    "Guardrail Hostname",
                    value=settings["guardrail_hostname"],
                    help="Example: https://www.us1.calypsoai.app",
                )
                guardrail_api_key = st.text_input(
                    "F5 Guardrail API key",
                    value="",
                    placeholder=(
                        "Configured for this session"
                        if settings["guardrail_api_key"]
                        else ""
                    ),
                    type="password",
                    help="Leave blank to keep the current session value.",
                )

            submitted = st.form_submit_button("Save settings")

        st.caption(
            "Values entered here are kept only for this browser session. "
            "Values already present in .env are loaded as defaults."
        )

        if submitted:
            updated_settings = {
                **settings,
                "model_provider": model_provider,
                "openai_model": openai_model.strip() or DEFAULT_OPENAI_MODEL,
                "openai_api_key": openai_api_key.strip() or settings["openai_api_key"],
                "openai_compatible_model": (
                    openai_compatible_model.strip() or DEFAULT_COMPATIBLE_MODEL
                ),
                "openai_compatible_base_url": (
                    openai_compatible_base_url.strip() or DEFAULT_COMPATIBLE_BASE_URL
                ),
                "openai_compatible_api_key": (
                    openai_compatible_api_key.strip()
                    or settings["openai_compatible_api_key"]
                ),
                "ollama_model": ollama_model.strip() or DEFAULT_OLLAMA_MODEL,
                "ollama_base_url": ollama_base_url.strip() or DEFAULT_OLLAMA_BASE_URL,
                "guardrail_api_key": (
                    guardrail_api_key.strip() or settings["guardrail_api_key"]
                ),
                "guardrail_hostname": guardrail_hostname.strip() or DEFAULT_GUARDRAIL_HOSTNAME,
            }
            settings = update_session_settings(st.session_state, updated_settings)
            st.success("Settings saved for this browser session")
            st.rerun()

        if st.button("Reset settings from .env"):
            settings = reset_session_settings(st.session_state, env_values)
            st.success("Session settings reset from .env")
            st.rerun()

    st.divider()

    st.caption(
        f"Active model: `{settings['model_provider']} / {get_selected_model(settings)}`"
    )


current_mode = (
    guardrail_enabled,
    guardrail_mode,
    settings["model_provider"],
    get_selected_model(settings),
)
if "last_mode" not in st.session_state:
    st.session_state.last_mode = current_mode
if st.session_state.last_mode != current_mode:
    st.session_state.messages = []
    st.session_state.last_mode = current_mode

if "messages" not in st.session_state:
    st.session_state.messages = []

for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])
        if "attachment" in msg:
            st.caption(format_attachment_caption(msg["attachment"]))
        if show_debug:
            render_debug_details(msg)

chat_submission = st.chat_input(
    "Enter your prompt...",
    accept_file=True,
    file_type=["pdf", "docx"],
    max_upload_size=10,
)

if chat_submission:
    if isinstance(chat_submission, str):
        prompt = chat_submission
        submitted_document = None
    else:
        prompt = chat_submission.text or ""
        submitted_document = chat_submission.files[0] if chat_submission.files else None

    if submitted_document is not None and not prompt.strip():
        prompt = "Please analyze the attached document."

    user_message = {"role": "user", "content": prompt}
    if submitted_document is not None:
        user_message["attachment"] = submitted_document.name

    st.session_state.messages.append(user_message)
    with st.chat_message("user"):
        st.markdown(prompt)
        if submitted_document is not None:
            st.caption(format_attachment_caption(submitted_document.name))

    extracted_document = None
    document_inspection_payload = prompt
    document_model_messages = None
    sensitive_debug_values = []

    if submitted_document is not None:
        try:
            extracted_document = extract_uploaded_document(submitted_document)
            document_inspection_payload = build_document_inspection_payload(
                prompt,
                extracted_document.filename,
                extracted_document.text,
            )
            document_model_messages = build_document_model_messages(
                prompt,
                extracted_document.filename,
                extracted_document.text,
            )
            sensitive_debug_values = [
                document_inspection_payload,
                extracted_document.text,
            ]
        except DocumentInspectionError as e:
            with st.chat_message("assistant"):
                st.error(str(e))
            st.session_state.messages.append(
                {
                    "role": "assistant",
                    "content": f"Document inspection failed: {e}",
                }
            )
            st.stop()
        except Exception as e:
            with st.chat_message("assistant"):
                st.error(f"Document parsing failed: {e}")
            st.session_state.messages.append(
                {
                    "role": "assistant",
                    "content": f"Document parsing failed: {e}",
                }
            )
            st.stop()

    if not guardrail_enabled:
        try:
            response_text = llm_chat(prompt, settings, messages=document_model_messages)
            st.session_state.messages.append({"role": "assistant", "content": response_text})
            with st.chat_message("assistant"):
                st.markdown(response_text)
        except Exception as e:
            with st.chat_message("assistant"):
                st.error(get_model_error_hint(e, settings) or f"Model call failed: {e}")
        st.stop()

    if guardrail_mode == "Inline":
        try:
            assistant_text, cai_json = cai_promptapi(
                document_inspection_payload,
                settings["guardrail_api_key"],
                settings["guardrail_prompt_api_url"],
            )
        except Exception as e:
            assistant_text, cai_json = None, {"error": str(e)}
        redacted_cai_json = redact_sensitive_debug_data(
            cai_json,
            sensitive_debug_values,
        )

        assistant_message = {
            "role": "assistant",
            "content": "⛔ Blocked / failed" if assistant_text is None else assistant_text,
            "cai_json": redacted_cai_json,
            "document": document_metadata(extracted_document),
        }
        with st.chat_message("assistant"):
            if assistant_text is None:
                st.error("Blocked / failed. See details in debug (enable 'Show debug details').")
            else:
                st.markdown(assistant_text)
            if show_debug:
                render_debug_details(assistant_message)

        st.session_state.messages.append(assistant_message)
        st.stop()

    try:
        cleared_in, scan_in_json = cai_scanapi(
            document_inspection_payload,
            settings["guardrail_api_key"],
            settings["guardrail_scan_url"],
        )
        if not cleared_in:
            redacted_scan_in_json = redact_sensitive_debug_data(
                scan_in_json,
                sensitive_debug_values,
            )
            with st.chat_message("assistant"):
                st.error("Prompt blocked due to policy.")
                assistant_message = {
                    "role": "assistant",
                    "content": "⛔ Prompt blocked due to policy.",
                    "scan_in": redacted_scan_in_json,
                    "scan_out": None,
                    "document": document_metadata(extracted_document),
                }
                if show_debug:
                    render_debug_details(assistant_message)
            st.session_state.messages.append(assistant_message)
            st.stop()

        response_text = llm_chat(prompt, settings, messages=document_model_messages)

        response_scan_payload = build_model_response_scan_payload(response_text)
        cleared_out, scan_out_json = cai_scanapi(
            response_scan_payload,
            settings["guardrail_api_key"],
            settings["guardrail_scan_url"],
        )
        if not cleared_out:
            redacted_scan_in_json = redact_sensitive_debug_data(
                scan_in_json,
                sensitive_debug_values,
            )
            redacted_scan_out_json = redact_sensitive_debug_data(
                scan_out_json,
                sensitive_debug_values,
            )
            with st.chat_message("assistant"):
                st.error("Response blocked due to policy.")
                assistant_message = {
                    "role": "assistant",
                    "content": "⛔ Response blocked due to policy.",
                    "scan_in": redacted_scan_in_json,
                    "scan_out": redacted_scan_out_json,
                    "document": document_metadata(extracted_document),
                }
                if show_debug:
                    render_debug_details(assistant_message)
            st.session_state.messages.append(assistant_message)
            st.stop()

        redacted_scan_in_json = redact_sensitive_debug_data(
            scan_in_json,
            sensitive_debug_values,
        )
        redacted_scan_out_json = redact_sensitive_debug_data(
            scan_out_json,
            sensitive_debug_values,
        )

        assistant_message = {
            "role": "assistant",
            "content": response_text,
            "scan_in": redacted_scan_in_json,
            "scan_out": redacted_scan_out_json,
            "document": document_metadata(extracted_document),
        }
        with st.chat_message("assistant"):
            st.markdown(response_text)
            if show_debug:
                render_debug_details(assistant_message)

        st.session_state.messages.append(assistant_message)

    except Exception as e:
        with st.chat_message("assistant"):
            st.error(get_model_error_hint(e, settings) or f"Out-of-band flow failed: {e}")
