from io import BytesIO

import pytest
from docx import Document

from f5_chatbot import (
    MAX_DOCUMENT_FILE_BYTES,
    MAX_EXTRACTED_TEXT_CHARS,
    DocumentInspectionError,
    extract_uploaded_document,
    build_document_inspection_payload,
    build_document_model_messages,
    extract_docx_text,
    validate_document_upload,
)


class FakeUpload:
    def __init__(self, name: str, size: int, data: bytes = b""):
        self.name = name
        self.size = size
        self._data = data

    def getvalue(self) -> bytes:
        return self._data


def test_validate_document_upload_accepts_pdf_under_limit():
    uploaded = FakeUpload("report.pdf", MAX_DOCUMENT_FILE_BYTES)

    assert validate_document_upload(uploaded) == ".pdf"


def test_validate_document_upload_rejects_unsupported_extension():
    uploaded = FakeUpload("report.txt", 100)

    with pytest.raises(DocumentInspectionError, match="Only PDF and DOCX"):
        validate_document_upload(uploaded)


def test_validate_document_upload_rejects_oversized_file():
    uploaded = FakeUpload("report.pdf", MAX_DOCUMENT_FILE_BYTES + 1)

    with pytest.raises(DocumentInspectionError, match="10 MB"):
        validate_document_upload(uploaded)


def test_build_document_inspection_payload_rejects_empty_text():
    with pytest.raises(DocumentInspectionError, match="No inspectable text"):
        build_document_inspection_payload("Summarize this", "empty.pdf", "   ")


def test_build_document_inspection_payload_rejects_oversized_text():
    oversized_text = "a" * (MAX_EXTRACTED_TEXT_CHARS + 1)

    with pytest.raises(DocumentInspectionError, match="too large"):
        build_document_inspection_payload("Summarize this", "large.pdf", oversized_text)


def test_build_document_inspection_payload_combines_prompt_and_document_text():
    payload = build_document_inspection_payload(
        "Summarize this",
        "report.pdf",
        "Document body",
    )

    assert "User prompt:" in payload
    assert "Summarize this" in payload
    assert "Attached document: report.pdf" in payload
    assert "Document body" in payload


def test_build_document_model_messages_marks_document_as_untrusted_context():
    messages = build_document_model_messages(
        "Summarize this",
        "report.docx",
        "Document body",
    )

    assert messages[0]["role"] == "system"
    assert "untrusted document content" in messages[0]["content"]
    assert messages[1]["role"] == "user"
    assert "report.docx" in messages[1]["content"]
    assert "Document body" in messages[1]["content"]
    assert "Summarize this" in messages[1]["content"]


def test_extract_docx_text_reads_paragraph_text():
    doc = Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    buffer = BytesIO()
    doc.save(buffer)

    text = extract_docx_text(buffer.getvalue())

    assert "First paragraph" in text
    assert "Second paragraph" in text


def test_extract_uploaded_document_rejects_actual_docx_bytes_over_limit():
    uploaded = FakeUpload(
        "large.docx",
        1,
        b"x" * (MAX_DOCUMENT_FILE_BYTES + 1),
    )

    with pytest.raises(DocumentInspectionError, match="10 MB"):
        extract_uploaded_document(uploaded)


def test_extract_uploaded_document_returns_docx_metadata():
    doc = Document()
    doc.add_paragraph("Paragraph content")
    buffer = BytesIO()
    doc.save(buffer)
    file_bytes = buffer.getvalue()
    uploaded = FakeUpload("notes.docx", 1, file_bytes)

    extracted = extract_uploaded_document(uploaded)

    assert extracted.filename == "notes.docx"
    assert extracted.extension == ".docx"
    assert extracted.size_bytes == len(file_bytes)
    assert "Paragraph content" in extracted.text
    assert extracted.char_count == len(extracted.text)


def test_extract_uploaded_document_rejects_empty_docx():
    doc = Document()
    buffer = BytesIO()
    doc.save(buffer)
    file_bytes = buffer.getvalue()
    uploaded = FakeUpload("empty.docx", len(file_bytes), file_bytes)

    with pytest.raises(DocumentInspectionError, match="No inspectable text"):
        extract_uploaded_document(uploaded)
